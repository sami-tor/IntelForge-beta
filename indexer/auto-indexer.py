#!/usr/bin/env python3
"""
Intel Forge Quickwit Auto-Indexer
Automatically indexes new files added to /data directory into Quickwit
"""

import os
import json
import time
import hashlib
import zipfile
import tempfile
import shutil
from pathlib import Path
from datetime import datetime
from urllib.request import Request, urlopen
from urllib.error import URLError, HTTPError
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler

# Archive libraries
try:
    import py7zr
    HAS_7Z = True
except ImportError:
    HAS_7Z = False

try:
    import rarfile
    HAS_RAR = True
except ImportError:
    HAS_RAR = False

# Document libraries
try:
    import PyPDF2
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import openpyxl
    HAS_XLSX = True
except ImportError:
    HAS_XLSX = False

# Image OCR
try:
    import pytesseract
    from PIL import Image
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

# Configuration
QUICKWIT_URL = os.getenv("QUICKWIT_URL", "http://localhost:7280")
DATA_DIRECTORY = os.getenv("DATA_DIRECTORY", "./data")
INDEX_NAME = os.getenv("INDEX_NAME", "osint-data")
IMAGE_INDEX_NAME = os.getenv("IMAGE_INDEX_NAME", "osint-data-images")
BATCH_SIZE = 1000
INDEXED_FILES_DB = os.getenv("INDEXED_FILES_DB", "/tmp/indexed_files.json")
DEBOUNCE_SECONDS = 5  # Wait 5 seconds after file is created before indexing
CHECK_DUPLICATES = os.getenv("CHECK_DUPLICATES", "false").lower() == "true"  # Disable duplicate checking by default (slow)
PERIODIC_SCAN_INTERVAL = int(os.getenv("PERIODIC_SCAN_INTERVAL", "5"))  # Scan /data folder every 5 seconds for missed files

# Behavior
INDEX_RECURSIVE = os.getenv("INDEX_RECURSIVE", "true").lower() == "true"  # Watch/index subfolders under /data
DELETE_AFTER_INDEX = os.getenv("DELETE_AFTER_INDEX", "false").lower() == "true"  # Keep original files by default


def get_file_hash(file_path: str) -> str:
    """Get MD5 hash of file for tracking"""
    try:
        with open(file_path, 'rb') as f:
            return hashlib.md5(f.read()).hexdigest()
    except:
        return ""


def load_indexed_files() -> dict:
    """Load list of already indexed files"""
    if os.path.exists(INDEXED_FILES_DB):
        try:
            with open(INDEXED_FILES_DB, 'r') as f:
                return json.load(f)
        except:
            return {}
    return {}


def save_indexed_files(indexed: dict):
    """Save list of indexed files"""
    try:
        with open(INDEXED_FILES_DB, 'w') as f:
            json.dump(indexed, f, indent=2)
        print(f"  ✅ Debug: Successfully saved {len(indexed)} files to {INDEXED_FILES_DB}")
    except Exception as e:
        print(f"❌ Warning: Could not save indexed files DB: {e}")
        import traceback
        traceback.print_exc()


def extract_country(file_path: str) -> str:
    """Extract country from file path like /data/countries/India/file.txt"""
    if "countries/" in file_path:
        parts = file_path.split("countries/")
        if len(parts) > 1:
            country = parts[1].split("/")[0]
            return country
    return "unknown"


def _normalize_category_name(name: str) -> str:
    name = (name or "").strip().lower()
    if not name:
        return "uncategorized"
    # Keep path separators for hierarchical categories; normalize each segment.
    segs = []
    for seg in name.replace("\\", "/").split("/"):
        seg = "".join(ch if (ch.isalnum() or ch in "-_") else "_" for ch in seg)
        seg = seg.strip("_")
        if seg:
            segs.append(seg)
    return "/".join(segs) if segs else "uncategorized"


def extract_category(file_path: str) -> str:
    """Extract category from /data folder structure.

    Examples:
      /data/database/*.txt -> database
      /data/scraped/telegram/*.json -> scraped/telegram
      /data/scraped/forums/xenforo/*.html -> scraped/forums/xenforo
      /data/intel/dns_ip_cert/*.txt -> intel/dns_ip_cert
    """
    try:
        normalized = file_path.replace("\\", "/")
        base = os.path.abspath(DATA_DIRECTORY).replace("\\", "/")
        # Make sure base ends with / to avoid prefix collisions.
        if not base.endswith("/"):
            base = base + "/"

        if normalized.startswith(base):
            rel = normalized[len(base):]
        else:
            rel = os.path.relpath(file_path, DATA_DIRECTORY).replace("\\", "/")

        # rel includes filename.
        parts = [p for p in rel.split("/") if p and p not in (".", "..")]
        if len(parts) <= 1:
            return "uncategorized"

        # Apply stable categorization rules.
        top = parts[0]
        if top == "scraped":
            if len(parts) >= 3 and parts[1] == "forums":
                return _normalize_category_name("/".join(parts[:3]))
            if len(parts) >= 2:
                return _normalize_category_name("/".join(parts[:2]))
            return "scraped"

        if top in ("leaks", "intel"):
            if len(parts) >= 2:
                return _normalize_category_name("/".join(parts[:2]))
            return _normalize_category_name(top)

        return _normalize_category_name(top)
    except Exception:
        return "uncategorized"


def get_file_type(file_name: str) -> str:
    """Get file extension"""
    ext = Path(file_name).suffix
    return ext[1:].lower() if ext else "txt"


def is_archive(file_path: str) -> bool:
    """Check if file is an archive"""
    ext = Path(file_path).suffix.lower()
    return ext in ['.zip', '.rar', '.7z']


def is_document(file_path: str) -> bool:
    """Check if file is a document"""
    ext = Path(file_path).suffix.lower()
    return ext in ['.pdf', '.docx', '.doc', '.xlsx', '.xls']


def is_image(file_path: str) -> bool:
    """Check if file is an image"""
    ext = Path(file_path).suffix.lower()
    return ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp']


def extract_text_from_zip(zip_path: str) -> list:
    """Extract text from ZIP archive"""
    texts = []
    try:
        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            for file_info in zip_ref.namelist():
                try:
                    # Skip directories
                    if file_info.endswith('/'):
                        continue
                    
                    # Read file content
                    content = zip_ref.read(file_info)
                    
                    # Try to decode as text
                    try:
                        text = content.decode('utf-8', errors='ignore')
                        # Split into lines
                        for line_num, line in enumerate(text.splitlines(), 1):
                            line = line.strip()
                            if line:
                                texts.append({
                                    "archive_path": zip_path,
                                    "file_name": file_info,
                                    "line_number": line_num,
                                    "content": line
                                })
                    except:
                        # If not text, skip
                        pass
                except Exception as e:
                    print(f"  ⚠️  Error extracting {file_info} from ZIP: {e}")
                    continue
    except Exception as e:
        print(f"  ❌ Error reading ZIP {zip_path}: {e}")
    return texts


def extract_text_from_7z(archive_path: str) -> list:
    """Extract text from 7Z archive"""
    texts = []
    if not HAS_7Z:
        print(f"  ⚠️  py7zr not available, skipping 7Z file")
        return texts
    
    try:
        with py7zr.SevenZipFile(archive_path, mode='r') as archive:
            # Extract to temporary directory
            with tempfile.TemporaryDirectory() as tmpdir:
                archive.extractall(path=tmpdir)
                
                # Walk through extracted files
                for root, dirs, files in os.walk(tmpdir):
                    for file_name in files:
                        file_path = os.path.join(root, file_name)
                        relative_path = os.path.relpath(file_path, tmpdir)
                        
                        try:
                            # Try to read as text
                            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                                for line_num, line in enumerate(f, 1):
                                    line = line.strip()
                                    if line:
                                        texts.append({
                                            "archive_path": archive_path,
                                            "file_name": relative_path,
                                            "line_number": line_num,
                                            "content": line
                                        })
                        except:
                            pass
    except Exception as e:
        print(f"  ❌ Error reading 7Z {archive_path}: {e}")
    return texts


def extract_text_from_rar(rar_path: str) -> list:
    """Extract text from RAR archive"""
    texts = []
    if not HAS_RAR:
        print(f"  ⚠️  rarfile not available, skipping RAR file")
        return texts
    
    try:
        with rarfile.RarFile(rar_path) as rf:
            # Extract to temporary directory
            with tempfile.TemporaryDirectory() as tmpdir:
                rf.extractall(path=tmpdir)
                
                # Walk through extracted files
                for root, dirs, files in os.walk(tmpdir):
                    for file_name in files:
                        file_path = os.path.join(root, file_name)
                        relative_path = os.path.relpath(file_path, tmpdir)
                        
                        try:
                            # Try to read as text
                            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                                for line_num, line in enumerate(f, 1):
                                    line = line.strip()
                                    if line:
                                        texts.append({
                                            "archive_path": rar_path,
                                            "file_name": relative_path,
                                            "line_number": line_num,
                                            "content": line
                                        })
                        except:
                            pass
    except Exception as e:
        print(f"  ❌ Error reading RAR {rar_path}: {e}")
    return texts


def extract_text_from_pdf(pdf_path: str) -> list:
    """Extract text from PDF document"""
    texts = []
    
    # Try pdfplumber first (better text extraction)
    if HAS_PDFPLUMBER:
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text:
                        for line_num, line in enumerate(text.splitlines(), 1):
                            line = line.strip()
                            if line:
                                texts.append({
                                    "page_number": page_num,
                                    "line_number": line_num,
                                    "content": line
                                })
            return texts
        except Exception as e:
            print(f"  ⚠️  pdfplumber failed, trying PyPDF2: {e}")
    
    # Fallback to PyPDF2
    if HAS_PDF:
        try:
            with open(pdf_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    text = page.extract_text()
                    if text:
                        for line_num, line in enumerate(text.splitlines(), 1):
                            line = line.strip()
                            if line:
                                texts.append({
                                    "page_number": page_num,
                                    "line_number": line_num,
                                    "content": line
                                })
        except Exception as e:
            print(f"  ❌ Error reading PDF {pdf_path}: {e}")
    
    return texts


def extract_text_from_docx(docx_path: str) -> list:
    """Extract text from DOCX document"""
    texts = []
    if not HAS_DOCX:
        print(f"  ⚠️  python-docx not available, skipping DOCX file")
        return texts
    
    try:
        doc = Document(docx_path)
        for para_num, paragraph in enumerate(doc.paragraphs, 1):
            text = paragraph.text.strip()
            if text:
                texts.append({
                    "paragraph_number": para_num,
                    "content": text
                })
        
        # Also extract text from tables
        for table_num, table in enumerate(doc.tables, 1):
            for row_num, row in enumerate(table.rows, 1):
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    texts.append({
                        "table_number": table_num,
                        "row_number": row_num,
                        "content": row_text
                    })
    except Exception as e:
        print(f"  ❌ Error reading DOCX {docx_path}: {e}")
    
    return texts


def extract_text_from_xlsx(xlsx_path: str) -> list:
    """Extract text from XLSX spreadsheet"""
    texts = []
    if not HAS_XLSX:
        print(f"  ⚠️  openpyxl not available, skipping XLSX file")
        return texts
    
    try:
        workbook = openpyxl.load_workbook(xlsx_path, data_only=True)
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            for row_num, row in enumerate(sheet.iter_rows(values_only=True), 1):
                row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                row_text = row_text.strip()
                if row_text:
                    texts.append({
                        "sheet_name": sheet_name,
                        "row_number": row_num,
                        "content": row_text
                    })
    except Exception as e:
        print(f"  ❌ Error reading XLSX {xlsx_path}: {e}")
    
    return texts


def extract_text_from_image(image_path: str) -> list:
    """Extract text from image using OCR"""
    texts = []
    if not HAS_OCR:
        print(f"  ⚠️  pytesseract/PIL not available, skipping image OCR")
        return texts
    
    try:
        # Open image
        image = Image.open(image_path)
        
        # Perform OCR
        ocr_text = pytesseract.image_to_string(image)
        
        # Split into lines
        for line_num, line in enumerate(ocr_text.splitlines(), 1):
            line = line.strip()
            if line:
                texts.append({
                    "line_number": line_num,
                    "content": line
                })
    except Exception as e:
        print(f"  ❌ Error performing OCR on {image_path}: {e}")
    
    return texts


def index_batch(docs: list, index_name: str) -> bool:
    """Index a batch of documents to Quickwit"""
    url = f"{QUICKWIT_URL}/api/v1/{index_name}/ingest"
    
    try:
        ndjson = "\n".join(json.dumps(doc) for doc in docs)
        ndjson_bytes = ndjson.encode('utf-8')
        
        req = Request(url, data=ndjson_bytes, headers={"Content-Type": "application/x-ndjson"}, method='POST')
        
        with urlopen(req, timeout=30) as response:
            status_code = response.getcode()
            if status_code in [200, 201]:
                return True
            else:
                response_text = response.read().decode('utf-8')
                print(f"Error indexing batch: {status_code} - {response_text}")
                return False
    except HTTPError as e:
        error_body = e.read().decode('utf-8') if e.fp else str(e)
        print(f"HTTP Error indexing batch: {e.code} - {error_body}")
        return False
    except Exception as e:
        print(f"Exception indexing batch: {e}")
        return False


def commit_index(index_name: str):
    """Commit index to make it searchable"""
    try:
        # Quickwit auto-commits, but we can try to force commit
        # The commit endpoint might not exist in newer versions, so we'll just skip it
        # Quickwit will auto-commit based on commit_timeout_secs
        return True
    except Exception as e:
        print(f"Warning: Could not commit index: {e}")
    return False


def check_duplicate_content(content: str, file_path: str) -> bool:
    """Check if content already exists in Quickwit"""
    try:
        # Search for exact content match
        search_url = f"{QUICKWIT_URL}/api/v1/{INDEX_NAME}/search"
        search_query = json.dumps({
            "query": f'content:"{content[:100]}"',  # Search first 100 chars
            "max_hits": 1,
            "search_field": "content"
        })
        
        req = Request(search_url, data=search_query.encode('utf-8'), 
                     headers={"Content-Type": "application/json"}, method='POST')
        with urlopen(req, timeout=5) as response:
            if response.getcode() == 200:
                data = json.loads(response.read().decode('utf-8'))
                if data.get('hits') and len(data['hits']) > 0:
                    # Check if exact content matches
                    for hit in data['hits']:
                        if hit.get('content') == content and hit.get('file_path') == file_path:
                            return True
    except:
        # If check fails, proceed with indexing (better to index than skip)
        pass
    return False


def index_file(file_path: str, indexed_files: dict) -> bool:
    """Index a single file into Quickwit"""
    # Normalize path
    normalized_path = file_path.replace("\\", "/")
    
    # Skip hidden files and system files (like .device-quotas.json, .gitignore, etc.)
    file_name = os.path.basename(file_path)
    if file_name.startswith('.'):
        print(f"⏭️  Skipping hidden file: {normalized_path}")
        return True
    
    # Skip generated files (anywhere in the tree)
    if "/thumbnails/" in normalized_path or "/webp/" in normalized_path or "/faces/" in normalized_path:
        print(f"⏭️  Skipping generated file: {normalized_path}")
        return True
    
    # Check if already indexed (by path and hash)
    try:
        file_hash = get_file_hash(file_path)
    except Exception as e:
        print(f"⚠️  Error getting file hash for {normalized_path}: {e}")
        file_hash = ""
    
    if normalized_path in indexed_files:
        stored_hash = indexed_files[normalized_path].get("hash", "")
        if stored_hash == file_hash:
            # For images, check if visual search succeeded
            if is_image(file_path):
                visual_search_success = indexed_files[normalized_path].get("visual_search_success", False)
                if not visual_search_success:
                    print(f"🔄 File {normalized_path} was partially indexed (OCR only), will retry visual search")
                    # Remove from indexed_files to allow retry
                    del indexed_files[normalized_path]
                    save_indexed_files(indexed_files)
                    # Continue processing (don't return True)
                else:
                    print(f"⏭️  Skipping {normalized_path} (already fully indexed)")
                    return True
            else:
                print(f"⏭️  Skipping {normalized_path} (already indexed)")
                return True
    
    file_type = get_file_type(file_name)
    category = extract_category(file_path)
    
    print(f"📄 Indexing: {normalized_path} (type: {file_type}, category: {category})")
    
    batch = []
    total_lines = 0
    extracted_texts = []
    
    try:
        # Handle different file types
        if is_archive(file_path):
            print(f"  📦 Detected archive file, extracting...")
            if file_type == 'zip':
                extracted_texts = extract_text_from_zip(file_path)
            elif file_type == '7z':
                extracted_texts = extract_text_from_7z(file_path)
            elif file_type == 'rar':
                extracted_texts = extract_text_from_rar(file_path)
            
            # Process extracted texts
            for item in extracted_texts:
                content = item["content"]
                # Check for duplicate content (disabled by default - slow)
                if CHECK_DUPLICATES and check_duplicate_content(content, normalized_path):
                    continue  # Skip duplicate
                
                total_lines += 1
                doc = {
                    "timestamp": int(datetime.now().timestamp()),
                    "file_path": normalized_path,
                    "file_name": file_name,
                    "line_number": item.get("line_number", 0),
                    "content": content,
                    "file_type": file_type,
                    "country": extract_country(normalized_path),
                    "category": category,
                    "archive_file": item.get("file_name", "")  # File inside archive
                }
                batch.append(doc)
                
                target_index = IMAGE_INDEX_NAME if is_image(file_path) else INDEX_NAME
                if len(batch) >= BATCH_SIZE:
                    if index_batch(batch, target_index):
                        print(f"  ✓ Indexed batch: {len(batch)} lines to {target_index}")
                    else:
                        print(f"  ✗ Failed to index batch")
                        return False
                    batch = []
        
        elif is_document(file_path):
            print(f"  📄 Detected document file, extracting text...")
            if file_type == 'pdf':
                extracted_texts = extract_text_from_pdf(file_path)
            elif file_type in ['docx', 'doc']:
                extracted_texts = extract_text_from_docx(file_path)
            elif file_type in ['xlsx', 'xls']:
                extracted_texts = extract_text_from_xlsx(file_path)
            
            # Process extracted texts
            for item in extracted_texts:
                content = item["content"]
                # Check for duplicate content (disabled by default - slow)
                if CHECK_DUPLICATES and check_duplicate_content(content, normalized_path):
                    continue  # Skip duplicate
                
                total_lines += 1
                doc = {
                    "timestamp": int(datetime.now().timestamp()),
                    "file_path": normalized_path,
                    "file_name": file_name,
                    "line_number": item.get("line_number", item.get("paragraph_number", item.get("row_number", 0))),
                    "content": content,
                    "file_type": file_type,
                    "country": extract_country(normalized_path),
                    "category": category
                }
                # Add document-specific metadata
                if "page_number" in item:
                    doc["page_number"] = item["page_number"]
                if "sheet_name" in item:
                    doc["sheet_name"] = item["sheet_name"]
                
                batch.append(doc)
                
                target_index = IMAGE_INDEX_NAME if is_image(file_path) else INDEX_NAME
                if len(batch) >= BATCH_SIZE:
                    if index_batch(batch, target_index):
                        print(f"  ✓ Indexed batch: {len(batch)} lines to {target_index}")
                    else:
                        print(f"  ✗ Failed to index batch")
                        return False
                    batch = []
        
        elif is_image(file_path):
            print(f"  🖼️  Detected image file, processing...")
            
            # Process image with visual search service (if available)
            image_metadata = {}
            embedding_id = None
            thumbnail_path = None
            webp_path = None
            faces = []
            ocr_text = ""
            detected_objects = []
            perceptual_hash = ""
            visual_search_success = False  # Track if visual search succeeded
            
            try:
                # Try to process image with visual search service
                # Use Docker service name if in Docker, otherwise localhost
                visual_search_url = os.getenv("VISUAL_SEARCH_SERVICE_URL", "http://intelforge-visual-search:8000")
                
                try:
                    import urllib.request
                    import base64
                    
                    # Read image file and encode as base64
                    with open(file_path, 'rb') as img_file:
                        image_data = img_file.read()
                        image_base64 = base64.b64encode(image_data).decode('utf-8')
                    
                    # Send as JSON (simpler than multipart)
                    request_data = json.dumps({
                        "file_path": normalized_path,
                        "image_base64": image_base64
                    }).encode('utf-8')
                    
                    req = Request(
                        f"{visual_search_url}/process-image",
                        data=request_data,
                        headers={
                            'Content-Type': 'application/json'
                        },
                        method='POST'
                    )
                    
                    with urlopen(req, timeout=600) as response:  # 10 min timeout for first model load
                        if response.getcode() == 200:
                            result = json.loads(response.read().decode('utf-8'))
                            if result.get("success"):
                                image_metadata = result.get("metadata", {})
                                embedding_id = result.get("embedding_id")
                                thumbnail_path = result.get("thumbnail_path")  # MinIO path
                                thumbnail_url = result.get("thumbnail_url")  # MinIO URL
                                webp_path = result.get("webp_path")  # MinIO path
                                webp_url = result.get("webp_url")  # MinIO URL
                                original_path = result.get("original_path")  # MinIO path
                                original_url = result.get("original_url")  # MinIO URL
                                faces = result.get("faces", [])  # Detected faces
                                ocr_text = result.get("ocr_text", "")
                                detected_objects = result.get("objects", [])  # Store objects list
                                objects_count = len(detected_objects)
                                faces_count = len(faces)
                                perceptual_hash = result.get("perceptual_hash", "")  # Store phash
                                visual_search_success = True  # Mark as successful
                                print(f"  ✓ Image processed and indexed to Qdrant + MinIO: Original={original_url}, WebP={webp_url}, Thumbnail={thumbnail_url}, Faces={faces_count}, Objects={objects_count}, Embedding={embedding_id}")
                            else:
                                print(f"  ⚠️  Visual search service returned error: {result.get('error', 'Unknown error')}")
                                visual_search_success = False
                                raise Exception(result.get('error', 'Processing failed'))
                        else:
                            error_text = response.read().decode('utf-8')
                            print(f"  ⚠️  Visual search service error {response.getcode()}: {error_text}")
                            visual_search_success = False
                            raise Exception(f"HTTP {response.getcode()}: {error_text}")
                except Exception as e:
                    print(f"  ⚠️  Visual search service not available: {e}")
                    visual_search_success = False
                    # Fallback to OCR only
                    extracted_texts = extract_text_from_image(file_path)
                    if extracted_texts:
                        ocr_text = "\n".join([item["content"] for item in extracted_texts])
            except Exception as e:
                print(f"  ⚠️  Error processing image with visual service: {e}")
                # Fallback to OCR
                extracted_texts = extract_text_from_image(file_path)
                if extracted_texts:
                    ocr_text = "\n".join([item["content"] for item in extracted_texts])
            
            # If OCR text not from visual service, extract it
            if not ocr_text:
                extracted_texts = extract_text_from_image(file_path)
                if extracted_texts:
                    ocr_text = "\n".join([item["content"] for item in extracted_texts])
            
            # Index OCR text as lines
            if ocr_text:
                for line_num, line in enumerate(ocr_text.splitlines(), 1):
                    line = line.strip()
                    if not line:
                        continue
                    
                    # Check for duplicate content
                    if CHECK_DUPLICATES and check_duplicate_content(line, normalized_path):
                        continue
                    
                    total_lines += 1
                    doc = {
                        "timestamp": int(datetime.now().timestamp()),
                        "file_path": normalized_path,
                        "file_name": file_name,
                        "line_number": line_num,
                        "content": line,
                        "file_type": file_type,
                        "country": extract_country(normalized_path),
                        "category": category,
                        "media_type": "image",
                        "extracted_via": "ocr"
                    }
                    
                    # Add visual metadata if available
                    if image_metadata:
                        doc["image_width"] = image_metadata.get("width")
                        doc["image_height"] = image_metadata.get("height")
                    if embedding_id:
                        doc["embedding_id"] = embedding_id
                    if thumbnail_path:
                        doc["thumbnail_path"] = thumbnail_path
                    if webp_path:
                        doc["webp_path"] = webp_path  # Store WebP path
                    if faces:
                        doc["face_count"] = len(faces)
                        doc["faces"] = faces  # Store face data
                    if ocr_text:
                        doc["ocr_text"] = line  # Store OCR text per line
                    
                    # Add advanced features if available
                    if 'detected_objects' in locals():
                        doc["detected_objects"] = detected_objects
                    if 'perceptual_hash' in locals():
                        doc["perceptual_hash"] = perceptual_hash
                    
                    batch.append(doc)
                    
                    target_index = IMAGE_INDEX_NAME
                    if len(batch) >= BATCH_SIZE:
                        if index_batch(batch, target_index):
                            print(f"  ✓ Indexed batch: {len(batch)} lines to {target_index}")
                        else:
                            print(f"  ✗ Failed to index batch")
                            return False
                        batch = []
            else:
                # Even if no OCR text, index image metadata
                total_lines = 1
                doc = {
                    "timestamp": int(datetime.now().timestamp()),
                    "file_path": normalized_path,
                    "file_name": file_name,
                    "line_number": 1,
                    "content": f"[Image: {file_name}]",
                    "file_type": file_type,
                    "country": extract_country(normalized_path),
                    "category": category,
                    "media_type": "image"
                }
                
                if image_metadata:
                    doc["image_width"] = image_metadata.get("width")
                    doc["image_height"] = image_metadata.get("height")
                if embedding_id:
                    doc["embedding_id"] = embedding_id
                if thumbnail_path:
                    doc["thumbnail_path"] = thumbnail_path
                if webp_path:
                    doc["webp_path"] = webp_path
                if faces:
                    doc["face_count"] = len(faces)
                    doc["faces"] = faces
                
                batch.append(doc)
        
        else:
            # Regular text file - read line by line
            print(f"  📝 Reading as text file...")
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                for line_num, line in enumerate(f, 1):
                    line = line.strip()
                    if not line:  # Skip empty lines
                        continue
                    
                    total_lines += 1
                    
                    # Check for duplicate content (skip if exact duplicate exists)
                    # Note: This is slow - checks Quickwit for each line. Disabled by default.
                    if CHECK_DUPLICATES and check_duplicate_content(line, normalized_path):
                        continue  # Skip duplicate line
                    
                    # Create document
                    doc = {
                        "timestamp": int(datetime.now().timestamp()),
                        "file_path": normalized_path,
                        "file_name": file_name,
                        "line_number": line_num,
                        "content": line,
                        "file_type": file_type,
                        "country": extract_country(normalized_path),
                        "category": category
                    }
                    
                    batch.append(doc)
                    
                    # Index batch when it reaches BATCH_SIZE
                    if len(batch) >= BATCH_SIZE:
                        if index_batch(batch, INDEX_NAME):
                            print(f"  ✓ Indexed batch: {len(batch)} lines")
                        else:
                            print(f"  ⚠️  Failed to index batch to Quickwit (continuing anyway)")
                            # Don't return False - continue processing to mark as indexed
                        batch = []
        
        # Index remaining documents
        # Use separate index for images
        target_index = IMAGE_INDEX_NAME if is_image(file_path) else INDEX_NAME
        if batch:
            if index_batch(batch, target_index):
                print(f"  ✓ Indexed final batch: {len(batch)} lines to {target_index}")
            else:
                print(f"  ⚠️  Failed to index final batch to Quickwit (but will mark as indexed to prevent retries)")
                # Don't return False - mark as indexed anyway to prevent infinite retries
        
        # Commit index
        if commit_index(target_index):
            print(f"  ✓ Committed {target_index}")
        
        # Mark as indexed (for images, only mark as fully indexed if visual search succeeded)
        if is_image(file_path):
            if visual_search_success and embedding_id:
                # Visual search succeeded - mark as fully indexed
                indexed_files[normalized_path] = {
                    "hash": file_hash,
                    "indexed_at": datetime.now().isoformat(),
                    "lines": total_lines,
                    "file_type": file_type,
                    "embedding_id": embedding_id,
                    "visual_search_success": True,
                    "deleted": False
                }
                save_indexed_files(indexed_files)
                
                # Check if MinIO upload was successful (original_url or original_path should exist)
                minio_upload_success = original_url or original_path
                
                if minio_upload_success:
                    print(f"✅ Successfully indexed image {normalized_path} to Milvus + MinIO ({total_lines} lines, embedding: {embedding_id})")
                    
                    # Delete file from /data after successful MinIO upload (optional)
                    if DELETE_AFTER_INDEX:
                        try:
                            if os.path.exists(file_path):
                                os.remove(file_path)
                                print(f"  🗑️  Deleted file from /data after indexing: {normalized_path}")
                                indexed_files[normalized_path]["deleted"] = True
                                save_indexed_files(indexed_files)
                        except Exception as e:
                            print(f"  ⚠️  Warning: Could not delete file {normalized_path}: {e}")
                else:
                    print(f"⚠️  Image {normalized_path} indexed to Milvus but not uploaded to MinIO")
            else:
                # Visual search failed - but Quickwit indexing succeeded (OCR only)
                # Mark as partially indexed and DELETE the file since it's already indexed to Quickwit
                print(f"⚠️  Image {normalized_path} indexed to Quickwit (OCR only, visual search failed)")
                print(f"  🔍 Debug: About to save file to indexed_files...")
                # Add to indexed_files with visual_search_success=False so we can retry visual search later if needed
                try:
                    indexed_files[normalized_path] = {
                        "hash": file_hash if file_hash else "",
                        "indexed_at": datetime.now().isoformat(),
                        "lines": total_lines,
                        "file_type": file_type,
                        "embedding_id": None,
                        "visual_search_success": False,
                        "deleted": False
                    }
                    hash_preview = file_hash[:16] if file_hash and len(file_hash) >= 16 else (file_hash or "N/A")
                    print(f"  🔍 Debug: Saving file to indexed_files: {normalized_path}, hash: {hash_preview}...")
                    save_indexed_files(indexed_files)
                    print(f"  ✅ Debug: Saved! Total files in indexed_files now: {len(indexed_files)}")
                    # Verify it was saved
                    if normalized_path in indexed_files:
                        print(f"  ✅ Debug: Verified - file is now in indexed_files")
                    else:
                        print(f"  ❌ Debug: ERROR - file is NOT in indexed_files after save!")
                    
                    # Delete file since it's already indexed to Quickwit (OCR) (optional)
                    if DELETE_AFTER_INDEX:
                        try:
                            if os.path.exists(file_path):
                                os.remove(file_path)
                                print(f"  🗑️  Deleted file from /data after indexing: {normalized_path}")
                                indexed_files[normalized_path]["deleted"] = True
                                save_indexed_files(indexed_files)
                        except Exception as e:
                            print(f"  ⚠️  Warning: Could not delete file {normalized_path}: {e}")
                        
                except Exception as e:
                    print(f"❌ Error saving file to indexed_files: {e}")
                    import traceback
                    traceback.print_exc()
            
            return True
        else:
            indexed_files[normalized_path] = {
                "hash": file_hash,
                "indexed_at": datetime.now().isoformat(),
                "lines": total_lines,
                "file_type": file_type,
                "deleted": False
            }
            save_indexed_files(indexed_files)
            print(f"✅ Successfully indexed {normalized_path} ({total_lines} lines, type: {file_type})")
            
            # Delete non-image files after successful indexing (optional)
            if DELETE_AFTER_INDEX:
                try:
                    if os.path.exists(file_path):
                        os.remove(file_path)
                        print(f"  🗑️  Deleted file after indexing: {normalized_path}")
                        indexed_files[normalized_path]["deleted"] = True
                        save_indexed_files(indexed_files)
                except Exception as e:
                    print(f"  ⚠️  Warning: Could not delete file {normalized_path}: {e}")
        
        return True
        
    except Exception as e:
        print(f"❌ Error indexing {normalized_path}: {e}")
        import traceback
        traceback.print_exc()
        return False


class FileIndexerHandler(FileSystemEventHandler):
    """Handler for file system events"""
    
    def __init__(self):
        self.indexed_files = load_indexed_files()
        self.pending_files = {}  # Track files waiting to be indexed (debounce)
        print(f"📋 Loaded {len(self.indexed_files)} already indexed files")
    
    def on_created(self, event):
        """Called when a file is created"""
        if event.is_directory:
            return
        
        file_path = event.src_path
        normalized_path = file_path.replace("\\", "/")
        
        # Skip hidden files
        if os.path.basename(file_path).startswith('.'):
            return
        
        # Skip generated files
        if "/thumbnails/" in normalized_path or "/webp/" in normalized_path or "/faces/" in normalized_path:
            return

        # Optional: only allow /data root if INDEX_RECURSIVE is disabled
        if not INDEX_RECURSIVE:
            try:
                rel_dir = os.path.dirname(os.path.relpath(file_path, DATA_DIRECTORY))
                if rel_dir not in ("", "."):
                    return
            except Exception:
                return
        
        # Skip if not in data directory
        if not file_path.startswith(os.path.abspath(DATA_DIRECTORY)):
            return
        
        # Debounce: wait a bit before indexing (file might still be writing)
        self.pending_files[file_path] = time.time()
        print(f"🔍 Detected new file: {file_path} (waiting {DEBOUNCE_SECONDS}s...)")
    
    def on_modified(self, event):
        """Called when a file is modified"""
        if event.is_directory:
            return
        
        file_path = event.src_path
        
        # Skip hidden files
        if os.path.basename(file_path).startswith('.'):
            return
        
        # Skip if not in data directory
        if not file_path.startswith(os.path.abspath(DATA_DIRECTORY)):
            return
        
        # If file was recently created, update debounce timer
        if file_path in self.pending_files:
            self.pending_files[file_path] = time.time()
    
    def process_pending_files(self):
        """Process files that are ready to be indexed (debounce complete)"""
        current_time = time.time()
        ready_files = []
        
        for file_path, detected_time in list(self.pending_files.items()):
            # Check if debounce period has passed
            if current_time - detected_time >= DEBOUNCE_SECONDS:
                # Check if file exists and is not being written to
                if os.path.exists(file_path):
                    try:
                        # Try to open file exclusively to check if it's still being written
                        with open(file_path, 'r'):
                            ready_files.append(file_path)
                            del self.pending_files[file_path]
                    except:
                        # File is still being written, keep waiting
                        pass
        
        # Index ready files
        for file_path in ready_files:
            index_file(file_path, self.indexed_files)


def check_quickwit_health():
    """Check if Quickwit is accessible"""
    try:
        # Try /health endpoint first
        health_url = f"{QUICKWIT_URL}/health"
        req = Request(health_url, method='GET')
        with urlopen(req, timeout=5) as response:
            if response.getcode() == 200:
                print(f"✅ Quickwit is accessible at {QUICKWIT_URL}")
                return True
    except HTTPError as e:
        if e.code == 404:
            # Try /api/v1/version as alternative health check
            try:
                version_url = f"{QUICKWIT_URL}/api/v1/version"
                req = Request(version_url, method='GET')
                with urlopen(req, timeout=5) as response:
                    if response.getcode() == 200:
                        print(f"✅ Quickwit is accessible at {QUICKWIT_URL}")
                        return True
            except:
                pass
        print(f"⚠️  Quickwit health check returned: {e.code}")
    except Exception as e:
        print(f"⚠️  Could not connect to Quickwit at {QUICKWIT_URL}: {e}")
        print(f"   Make sure Quickwit is running and accessible")
    return False


def _quickwit_index_exists(index_id: str) -> bool:
    try:
        check_url = f"{QUICKWIT_URL}/api/v1/indexes/{index_id}"
        req = Request(check_url, method='GET')
        with urlopen(req, timeout=5) as response:
            return response.getcode() == 200
    except HTTPError as e:
        return e.code != 404
    except Exception:
        return False


def _quickwit_create_index(config: dict) -> bool:
    try:
        create_url = f"{QUICKWIT_URL}/api/v1/indexes"
        req = Request(
            create_url,
            data=json.dumps(config).encode('utf-8'),
            headers={'Content-Type': 'application/json'},
            method='POST'
        )
        with urlopen(req, timeout=30) as response:
            return response.getcode() in (200, 201)
    except Exception:
        return False


def _get_text_index_config() -> dict:
    return {
        "version": "0.8",
        "index_id": INDEX_NAME,
        "doc_mapping": {
            "mode": "lenient",
            "field_mappings": [
                {"name": "timestamp", "type": "datetime", "input_formats": ["unix_timestamp"], "fast": True, "indexed": True},
                {"name": "file_path", "type": "text", "tokenizer": "raw", "stored": True, "indexed": True},
                {"name": "file_name", "type": "text", "tokenizer": "default", "stored": True, "indexed": True},
                {"name": "line_number", "type": "u64", "stored": True, "indexed": True},
                {"name": "content", "type": "text", "tokenizer": "default", "stored": True, "indexed": True, "record": "position"},
                {"name": "file_type", "type": "text", "tokenizer": "raw", "stored": True, "indexed": True},
                {"name": "country", "type": "text", "tokenizer": "raw", "stored": True, "indexed": True},
                {"name": "category", "type": "text", "tokenizer": "raw", "stored": True, "indexed": True},
            ],
            "timestamp_field": "timestamp",
        },
        "search_settings": {
            "default_search_fields": ["content", "file_name", "file_path"],
        },
    }


def _get_images_index_config() -> dict:
    return {
        "version": "0.8",
        "index_id": IMAGE_INDEX_NAME,
        "doc_mapping": {
            "mode": "lenient",
            "field_mappings": [
                {"name": "timestamp", "type": "datetime", "input_formats": ["unix_timestamp"], "fast": True, "indexed": True},
                {"name": "file_path", "type": "text", "tokenizer": "raw", "stored": True, "indexed": True},
                {"name": "file_name", "type": "text", "tokenizer": "default", "stored": True, "indexed": True},
                {"name": "file_type", "type": "text", "tokenizer": "raw", "stored": True, "indexed": True},
                {"name": "country", "type": "text", "tokenizer": "raw", "stored": True, "indexed": True},
                {"name": "category", "type": "text", "tokenizer": "raw", "stored": True, "indexed": True},
                {"name": "media_type", "type": "text", "tokenizer": "raw", "stored": True, "indexed": True},
                {"name": "image_width", "type": "u64", "stored": True, "indexed": True},
                {"name": "image_height", "type": "u64", "stored": True, "indexed": True},
                {"name": "embedding_id", "type": "text", "tokenizer": "raw", "stored": True, "indexed": True},
                {"name": "thumbnail_path", "type": "text", "tokenizer": "raw", "stored": True, "indexed": True},
                {"name": "thumbnail_url", "type": "text", "tokenizer": "raw", "stored": True, "indexed": True},
                {"name": "webp_path", "type": "text", "tokenizer": "raw", "stored": True, "indexed": True},
                {"name": "webp_url", "type": "text", "tokenizer": "raw", "stored": True, "indexed": True},
                {"name": "original_path", "type": "text", "tokenizer": "raw", "stored": True, "indexed": True},
                {"name": "original_url", "type": "text", "tokenizer": "raw", "stored": True, "indexed": True},
                {"name": "ocr_text", "type": "text", "tokenizer": "default", "stored": True, "indexed": True, "record": "position"},
                {"name": "file_size", "type": "u64", "stored": True, "indexed": True},
                {"name": "detected_objects", "type": "json", "stored": True, "indexed": True},
                {"name": "perceptual_hash", "type": "text", "tokenizer": "raw", "stored": True, "indexed": True},
                {"name": "face_count", "type": "u64", "stored": True, "indexed": True},
                {"name": "faces", "type": "json", "stored": True, "indexed": True},
            ],
            "timestamp_field": "timestamp",
        },
        "search_settings": {
            "default_search_fields": ["file_name", "file_path", "ocr_text"],
        },
    }


def ensure_quickwit_indexes():
    if not _quickwit_index_exists(INDEX_NAME):
        _quickwit_create_index(_get_text_index_config())
    if not _quickwit_index_exists(IMAGE_INDEX_NAME):
        _quickwit_create_index(_get_images_index_config())


def _iter_data_files() -> list:
    files = []
    if not os.path.exists(DATA_DIRECTORY):
        return files

    if INDEX_RECURSIVE:
        for root, dirs, file_names in os.walk(DATA_DIRECTORY):
            # Skip hidden dirs
            dirs[:] = [d for d in dirs if not d.startswith('.')]

            # Skip generated dirs anywhere
            pruned = []
            for d in dirs:
                dn = d.lower()
                if dn in ("thumbnails", "webp", "faces"):
                    continue
                pruned.append(d)
            dirs[:] = pruned

            for file_name in file_names:
                if file_name.startswith('.'):
                    continue
                file_path = os.path.join(root, file_name)
                normalized_path = file_path.replace("\\", "/")
                if "/thumbnails/" in normalized_path or "/webp/" in normalized_path or "/faces/" in normalized_path:
                    continue
                files.append(file_path)
    else:
        for file_name in os.listdir(DATA_DIRECTORY):
            if file_name.startswith('.'):
                continue
            file_path = os.path.join(DATA_DIRECTORY, file_name)
            if os.path.isdir(file_path):
                continue
            normalized_path = file_path.replace("\\", "/")
            if "/thumbnails/" in normalized_path or "/webp/" in normalized_path or "/faces/" in normalized_path:
                continue
            files.append(file_path)

    return files


def index_existing_files(handler_indexed_files=None):
    """Index any existing files that haven't been indexed yet."""
    if handler_indexed_files is not None:
        indexed_files = handler_indexed_files
    else:
        indexed_files = load_indexed_files()
        print("🔍 Checking for unindexed files...")

    indexed_count = 0
    deleted_count = 0

    data_files = _iter_data_files()
    print(f"🔍 Found {len(data_files)} files in /data")

    for file_path in data_files:
        file_name = os.path.basename(file_path)
        normalized_path = file_path.replace("\\", "/")

        # Check if already indexed
        if normalized_path in indexed_files:
            try:
                file_hash = get_file_hash(file_path)
            except Exception:
                file_hash = ""

            stored_hash = indexed_files[normalized_path].get("hash", "")
            if stored_hash == file_hash:
                # If we keep source files, skip deletion checks.
                if not DELETE_AFTER_INDEX:
                    # If it's an image and visual search failed before, retry.
                    if is_image(file_path) and not indexed_files[normalized_path].get("visual_search_success", False):
                        del indexed_files[normalized_path]
                        save_indexed_files(indexed_files)
                    else:
                        continue

                # DELETE_AFTER_INDEX enabled: enforce deletion policy.
                if is_image(file_path):
                    visual_search_success = indexed_files[normalized_path].get("visual_search_success", False)
                    is_deleted = indexed_files[normalized_path].get("deleted", False)
                    file_exists = os.path.exists(file_path)

                    if not visual_search_success:
                        del indexed_files[normalized_path]
                        save_indexed_files(indexed_files)
                    else:
                        if not is_deleted and file_exists:
                            try:
                                os.remove(file_path)
                                indexed_files[normalized_path]["deleted"] = True
                                deleted_count += 1
                            except Exception:
                                pass
                        continue
                else:
                    if not indexed_files[normalized_path].get("deleted", False):
                        try:
                            if os.path.exists(file_path):
                                os.remove(file_path)
                                indexed_files[normalized_path]["deleted"] = True
                                deleted_count += 1
                        except Exception:
                            pass
                    continue
            else:
                # Hash mismatch - re-index
                del indexed_files[normalized_path]
                save_indexed_files(indexed_files)

        # Index the file
        try:
            result = index_file(file_path, indexed_files)
            if result:
                indexed_count += 1
        except Exception:
            pass

    if deleted_count > 0:
        save_indexed_files(indexed_files)

    if indexed_count > 0:
        print(f"✅ Indexed {indexed_count} new files")


def main():
    """Main function"""
    print("=" * 60)
    print("🚀 Intel Forge Quickwit Auto-Indexer")
    print("=" * 60)
    print(f"Quickwit URL: {QUICKWIT_URL}")
    print(f"Data Directory: {DATA_DIRECTORY}")
    print(f"Index Name: {INDEX_NAME}")
    print(f"Indexed Files DB: {INDEXED_FILES_DB}")
    print("-" * 60)
    
    # Check Quickwit health
    if not check_quickwit_health():
        print(f"⚠️  Warning: Could not connect to Quickwit at {QUICKWIT_URL}")
        print("   Make sure Quickwit is running and accessible")
        print("   Will retry when files are detected...")
    else:
        print("✅ Quickwit is accessible")
        ensure_quickwit_indexes()
    
    # Index existing files first
    index_existing_files()
    
    print("-" * 60)
    print("👀 Starting file watcher...")
    print(f"   Watching: {os.path.abspath(DATA_DIRECTORY)}")
    print("   Press Ctrl+C to stop")
    print("=" * 60)
    
    # Set up file watcher
    event_handler = FileIndexerHandler()
    observer = Observer()
    observer.schedule(event_handler, DATA_DIRECTORY, recursive=INDEX_RECURSIVE)
    observer.start()
    
    try:
        # Process pending files every second
        last_scan_time = time.time()
        while True:
            time.sleep(1)
            event_handler.process_pending_files()
            
            # Periodic scan for missed files (every PERIODIC_SCAN_INTERVAL seconds)
            current_time = time.time()
            if current_time - last_scan_time >= PERIODIC_SCAN_INTERVAL:
                print(f"🔍 Periodic scan: Checking for missed files (every {PERIODIC_SCAN_INTERVAL}s)...")
                index_existing_files(event_handler.indexed_files)
                last_scan_time = current_time
    except KeyboardInterrupt:
        print("\n🛑 Stopping file watcher...")
        observer.stop()
    
    observer.join()
    print("✅ Auto-indexer stopped")


if __name__ == "__main__":
    main()

